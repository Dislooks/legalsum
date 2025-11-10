"""
Dokumentenleser für verschiedene Dateiformate (PDF, DOCX, TXT).

Dieses Modul bietet eine einheitliche Schnittstelle zum Einlesen
juristischer Dokumente in verschiedenen Formaten.
"""

import os
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from abc import ABC, abstractmethod
import logging

# PDF-Verarbeitung
try:
    import pikepdf
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX-Verarbeitung
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

# Bei Entwicklung, Testing, Debugging und Erweiterung wurde Claude Sonnet 4.5 genutzt
class DocumentReader(ABC):
    """Abstrakte Basisklasse für Dokumentenleser."""
    
    @abstractmethod
    def read(self, file_path: Path) -> Dict[str, str]:
        """
        Liest ein Dokument und gibt strukturierte Daten zurück.
        
        Args:
            file_path: Pfad zur Datei

        """
        pass
    
    def _extract_metadata(self, file_path: Path) -> Dict[str, str]:
        """Extrahiert grundlegende Dateimetadaten."""
        stat = file_path.stat()
        return {
            'filename': file_path.name,
            'file_size': stat.st_size,
            'file_extension': file_path.suffix,
            'absolute_path': str(file_path.absolute())
        }

# Bei Entwicklung, Testing, Debugging und Erweiterung wurde Claude Sonnet 4.5 genutzt
class PDFReader(DocumentReader):
    """
    PDF-Dokumentenleser mit Fallback-Strategie.
    """
    
    def __init__(self, use_ocr: bool = False):
        """
        Args:
            use_ocr: Aktiviert OCR für gescannte PDFs (erfordert tesseract)
        """
        if not PDF_AVAILABLE:
            raise ImportError(
                "PDF-Bibliotheken nicht installiert. "
                "Installiere mit: pip install pikepdf pdfplumber"
            )
        self.use_ocr = use_ocr
        
    def read(self, file_path: Path) -> Dict[str, str]:
        """Liest PDF-Dokument."""
        try:
            # Primäre Methode: pdfplumber (bessere Textextraktion)
            text = self._read_with_pdfplumber(file_path)
            method = "pdfplumber"
        except Exception as e:
            logger.warning(f"pdfplumber fehlgeschlagen: {e}. Verwende pikepdf.")
            # Fallback: pikepdf
            text = self._read_with_pikepdf(file_path)
            method = "pikepdf"
        
        metadata = self._extract_metadata(file_path)
        metadata['extraction_method'] = method
        metadata['page_count'] = self._get_page_count(file_path)
        
        return {
            'text': text,
            'metadata': metadata,
            'file_info': {
                'type': 'pdf',
                'encoding': 'utf-8'
            }
        }
    
    def _read_with_pdfplumber(self, file_path: Path) -> str:
        """Extrahiert Text mit pdfplumber."""
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        return '\n\n'.join(text_parts)
    
    def _read_with_pikepdf(self, file_path: Path) -> str:
        """Extrahiert Text mit pikepdf (Fallback)."""
        text_parts = []
        with open(file_path, 'rb') as file:
            pdf_reader = pikepdf.PdfReader(file)
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        
        return '\n\n'.join(text_parts)
    
    def _get_page_count(self, file_path: Path) -> int:
        """Ermittelt Seitenanzahl."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pikepdf.PdfReader(file)
                return len(pdf_reader.pages)
        except:
            return 0

# Bei Entwicklung, Testing, Debugging und Erweiterung wurde Claude Sonnet 4.5 genutzt
class DOCXReader(DocumentReader):
    """Microsoft Word DOCX-Dokumentenleser."""
    
    def __init__(self, preserve_formatting: bool = False):
        """
        Args:
            preserve_formatting: Erhält Formatierungsinformationen
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx nicht installiert. "
                "Installiere mit: pip install python-docx"
            )
        self.preserve_formatting = preserve_formatting
    
    def read(self, file_path: Path) -> Dict[str, str]:
        """Liest DOCX-Dokument."""
        doc = Document(file_path)
        
        # Extrahiere Text aus Absätzen
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        text = '\n\n'.join(paragraphs)
        
        # Extrahiere Text aus Tabellen
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    table_texts.append(row_text)
        
        if table_texts:
            text += '\n\n' + '\n'.join(table_texts)
        
        metadata = self._extract_metadata(file_path)
        metadata['paragraph_count'] = len(paragraphs)
        metadata['table_count'] = len(doc.tables)
        
        # Core Properties (falls vorhanden)
        core_props = doc.core_properties
        metadata['author'] = getattr(core_props, 'author', 'Unknown')
        metadata['title'] = getattr(core_props, 'title', '')
        
        return {
            'text': text,
            'metadata': metadata,
            'file_info': {
                'type': 'docx',
                'encoding': 'utf-8'
            }
        }

# Bei Entwicklung, Testing, Debugging und Erweiterung wurde Claude Sonnet 4.5 genutzt
class TXTReader(DocumentReader):
    """Einfacher Textdatei-Leser mit Encoding-Erkennung."""
    
    def __init__(self, encoding: Optional[str] = None):
        """
        Args:
            encoding: Zeichenkodierung (None = automatische Erkennung)
        """
        self.encoding = encoding
    
    def read(self, file_path: Path) -> Dict[str, str]:
        """Liest TXT-Datei."""
        # Encoding-Erkennung
        if self.encoding is None:
            detected_encoding = self._detect_encoding(file_path)
        else:
            detected_encoding = self.encoding
        
        # Datei einlesen
        try:
            with open(file_path, 'r', encoding=detected_encoding) as file:
                text = file.read()
        except UnicodeDecodeError:
            logger.warning(f"Encoding {detected_encoding} fehlgeschlagen. Verwende latin-1.")
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
            detected_encoding = 'latin-1'
        
        metadata = self._extract_metadata(file_path)
        metadata['line_count'] = text.count('\n') + 1
        
        return {
            'text': text,
            'metadata': metadata,
            'file_info': {
                'type': 'txt',
                'encoding': detected_encoding
            }
        }
    
    def _detect_encoding(self, file_path: Path) -> str:
        """
        Erkennt Dateienkodierung automatisch.
        
        Versucht gängige Encodings in dieser Reihenfolge:
        utf-8, utf-8-sig, latin-1
        """
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    file.read()
                return encoding
            except (UnicodeDecodeError, LookupError):
                continue
        
        logger.warning("Encoding-Erkennung fehlgeschlagen. Verwende utf-8.")
        return 'utf-8'

# Bei Entwicklung, Testing, Debugging und Erweiterung wurde Claude Sonnet 4.5 genutzt
class DocumentReaderFactory:
    """
    Factory-Klasse zur automatischen Auswahl des passenden Readers.

    """
    
    _readers = {
        '.pdf': PDFReader,
        '.docx': DOCXReader,
        '.doc': DOCXReader,  # .doc wird wie .docx behandelt
        '.txt': TXTReader
    }
    
    @classmethod
    def get_reader(cls, file_path: Path, **kwargs) -> DocumentReader:
        """
        Gibt den passenden Reader für eine Datei zurück.
        
        Args:
            file_path: Pfad zur Datei
            **kwargs: Zusätzliche Parameter für den Reader

        """
        if isinstance(file_path, str):
            file_path = Path(file_path)
        
        extension = file_path.suffix.lower()
        
        if extension not in cls._readers:
            supported = ', '.join(cls._readers.keys())
            raise ValueError(
                f"Dateityp '{extension}' nicht unterstützt. "
                f"Unterstützte Formate: {supported}"
            )
        
        reader_class = cls._readers[extension]
        return reader_class(**kwargs)
    
    @classmethod
    def read_document(cls, file_path: Path, **kwargs) -> Dict[str, str]:
        """
        Convenience-Methode zum direkten Einlesen eines Dokuments.
        
        Args:
            file_path: Pfad zur Datei
            **kwargs: Zusätzliche Parameter für den Reader
 
        """
        reader = cls.get_reader(file_path, **kwargs)
        return reader.read(file_path)

# Bei Entwicklung, Testing, Debugging und Erweiterung wurde Claude Sonnet 4.5 genutzt
def read_documents_from_directory(
    directory: Path,
    extensions: Optional[List[str]] = None,
    recursive: bool = False,
    **reader_kwargs
) -> List[Dict[str, str]]:
    """
    Liest alle Dokumente aus einem Verzeichnis.
    
    Args:
        directory: Quellverzeichnis
        extensions: Liste erlaubter Dateierweiterungen (z.B. ['.pdf', '.docx'])
        recursive: Rekursiv in Unterverzeichnisse
        **reader_kwargs: Parameter für die Reader
    
    """
    if isinstance(directory, str):
        directory = Path(directory)
    
    if not directory.exists():
        raise FileNotFoundError(f"Verzeichnis nicht gefunden: {directory}")
    
    # Standard: alle unterstützten Formate
    if extensions is None:
        extensions = ['.pdf', '.docx', '.doc', '.txt']
    
    # Dateien sammeln
    pattern = '**/*' if recursive else '*'
    all_files = directory.glob(pattern)
    
    documents = []
    for file_path in all_files:
        if file_path.is_file() and file_path.suffix.lower() in extensions:
            try:
                logger.info(f"Lese Dokument: {file_path.name}")
                doc = DocumentReaderFactory.read_document(file_path, **reader_kwargs)
                documents.append(doc)
            except Exception as e:
                logger.error(f"Fehler beim Lesen von {file_path.name}: {e}")
                continue
    
    logger.info(f"Erfolgreich {len(documents)} Dokumente eingelesen.")
    return documents
